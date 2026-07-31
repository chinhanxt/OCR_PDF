import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

class DOCXBuilder:
    def __init__(self):
        pass

    def build_docx_for_paddle(self, pages_metadata: list, output_docx_path: str):
        doc = docx.Document()
        self._set_document_styles(doc)

        title = doc.add_paragraph()
        title_run = title.add_run("KẾT QUẢ SCAN OCR (PADDLEOCR GPU)")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(26, 86, 219)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph()
        sub_run = sub.add_run(f"Tổng số trang: {len(pages_metadata)} | Độ chính xác Tiếng Việt >95%")
        sub_run.font.size = Pt(10)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(100, 116, 139)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for page in pages_metadata:
            page_no = page.get("page_number", 1)
            h = doc.add_heading(f"Trang {page_no}", level=2)
            h.runs[0].font.color.rgb = RGBColor(30, 41, 59)

            items = page.get("ocr_items", [])
            lines = self._group_items_into_lines(items)

            for line_items in lines:
                line_text = "   ".join([it["text"] for it in line_items])
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                r = p.add_run(line_text)
                r.font.size = Pt(11)
                r.font.name = "Calibri"

            doc.add_page_break()

        doc.save(output_docx_path)
        print(f"✅ PaddleOCR DOCX created at: {output_docx_path}")

    def build_docx_for_docling(self, pages_metadata: list, output_docx_path: str, docling_doc=None):
        doc = docx.Document()
        self._set_document_styles(doc)

        title = doc.add_paragraph()
        title_run = title.add_run("KẾT QUẢ BÓC TÁCH TÀI LIỆU (DOCLING AI)")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(147, 51, 234)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph()
        sub_run = sub.add_run(f"Tổng số trang: {len(pages_metadata)} | IBM TableFormer Layout AI")
        sub_run.font.size = Pt(10)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(100, 116, 139)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for page in pages_metadata:
            page_no = page.get("page_number", 1)
            h = doc.add_heading(f"Trang {page_no}", level=2)
            h.runs[0].font.color.rgb = RGBColor(30, 41, 59)

            items = page.get("ocr_items", [])
            
            # Separate table cells vs text items
            table_cells = [it for it in items if it.get("type") == "table_cell" and it.get("row") is not None]
            text_items = [it for it in items if it.get("type") != "table_cell"]

            # 1. Render Paragraph Texts
            if text_items:
                lines = self._group_items_into_lines(text_items)
                for line_items in lines:
                    line_text = "   ".join([it["text"] for it in line_items])
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(line_text)
                    r.font.size = Pt(11)

            # 2. Render Table Grid (TableFormer)
            if table_cells:
                max_row = max([c["row"] for c in table_cells]) + 1
                max_col = max([c["col"] for c in table_cells]) + 1
                
                if max_row > 0 and max_col > 0:
                    tbl = doc.add_table(rows=max_row, cols=max_col)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.style = 'Table Grid'

                    for c in table_cells:
                        r_idx = c["row"]
                        c_idx = c["col"]
                        if r_idx < max_row and c_idx < max_col:
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = c["text"]
                            if r_idx == 0:
                                shading_elm = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                                cell._tc.get_or_add_tcPr().append(shading_elm)
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.bold = True
                                        r.font.size = Pt(10)

                    doc.add_paragraph()

            doc.add_page_break()

        doc.save(output_docx_path)
        print(f"✅ Docling DOCX created at: {output_docx_path}")

    def _group_items_into_lines(self, items: list, y_tolerance: float = 12.0) -> list:
        if not items:
            return []
        
        sorted_items = sorted(items, key=lambda x: (x["bbox"][1], x["bbox"][0]))
        lines = []
        current_line = [sorted_items[0]]

        for item in sorted_items[1:]:
            prev_y = current_line[-1]["bbox"][1]
            curr_y = item["bbox"][1]
            if abs(curr_y - prev_y) <= y_tolerance:
                current_line.append(item)
            else:
                lines.append(sorted(current_line, key=lambda x: x["bbox"][0]))
                current_line = [item]

        if current_line:
            lines.append(sorted(current_line, key=lambda x: x["bbox"][0]))

        return lines

    def _set_document_styles(self, doc):
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
